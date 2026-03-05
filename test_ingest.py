import unittest
import docx
import os

from ingest import read_docx

class TestIngestion(unittest.TestCase):
    
    def setUp(self):
        """Тази функция се пуска ПРЕДИ всеки тест. Тук създаваме фалшив Word файл."""
        self.test_file = "dummy_test_meeting.docx"
        doc = docx.Document()
        doc.add_paragraph("Това е тестова среща.")
        doc.add_paragraph("Обсъждаме изкуствен интелект.")
        doc.save(self.test_file)

    def test_read_docx(self):
        """Това е същинският тест. Проверяваме дали функцията чете правилно."""
        extracted_text = read_docx(self.test_file)
        
        self.assertIn("Това е тестова среща.", extracted_text)
        self.assertIn("Обсъждаме изкуствен интелект.", extracted_text)
        
        print("\n✅ ТЕСТЪТ МИНА УСПЕШНО! Функцията read_docx работи перфектно.")

    def tearDown(self):
        """Тази функция се пуска СЛЕД теста. Тук изтриваме фалшивия файл."""
        if os.path.exists(self.test_file):
            os.remove(self.test_file)

if __name__ == '__main__':
    unittest.main()